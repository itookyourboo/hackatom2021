from docx import Document
from docx.shared import RGBColor

from util import DENSITY, PARAGRAPH_EVERY_N_SENTENCES


class GoogleCloudResultsAnalyzer:
    def __init__(self, results):
        self.results = results

    def average_confidence(self):
        return sum([result.alternatives[0].confidence for result in self.results]) / len(self.results)

    def get_debug_info(self):
        arr = ['========================']
        for result in self.results:
            arr.append(f'Transcription\t{[x.transcript for x in result.alternatives]}')
            arr.append(f'Confidence\t{[x.confidence for x in result.alternatives]}')
        arr.append('========================')

        return '\n'.join(arr)

    def get_text(self):
        return '\n'.join([result.alternatives[0].transcript.strip() for result in self.results])


    def get_formatted_text(self, text):
        sentences = text.split('. ')
        result = []
        for i, sentence in enumerate(sentences):
            result.append(sentence)
            if i % PARAGRAPH_EVERY_N_SENTENCES == PARAGRAPH_EVERY_N_SENTENCES - 1:
                result.append('.\n')
            else:
                result.append('. ')

        return ''.join(result)


class VoskResultsAnalyzer:
    def __init__(self, words, lang='ru'):
        self.words = words
        self.lang = lang

    def get_punctuation(self):
        density = DENSITY[self.lang]
        TOTAL_WORDS = len(self.words)

        self.words[0].value = self.words[0].value.capitalize()
        punctuation = {}
        pairs = []
        for i in range(1, len(self.words)):
            pairs.append((self.words[i].start - self.words[i - 1].end, i - 1))

        pairs.sort(key=lambda x: x[0])
        dots_count = int(TOTAL_WORDS * density['dot'])
        comma_count = int(TOTAL_WORDS * density['comma'])
        paragraph_count = int(TOTAL_WORDS * density['paragraph'])

        counter = 0

        for pair in pairs[-(comma_count + dots_count):]:
            counter += 1

            if counter <= comma_count:
                punctuation[pair[1]] = ','
            elif counter <= comma_count + dots_count - paragraph_count:
                punctuation[pair[1]] = '.'
            else:
                punctuation[pair[1]] = '.\n'

        return punctuation

    def get_confidence(self):
        return sum(word.conf for word in self.words) / len(self.words)

    def get_confidence_percent(self):
        confidence = self.get_confidence()
        return int(confidence * 10000) / 100

    def to_plain_text(self):
        punctuation = self.get_punctuation()

        lines = []
        line = []

        for i, word in enumerate(self.words):
            if i in punctuation:
                if '.' in punctuation[i]:
                    self.words[i + 1].value = self.words[i + 1].value.capitalize()
                line.append(word.value + punctuation[i].rstrip('\n'))

                if '\n' in punctuation[i]:
                    lines.append(' '.join(line))
                    line.clear()
            else:
                line.append(word.value)

        return '\n'.join(lines)

    def to_docx(self, file_name):
        document = Document()

        punctuation = self.get_punctuation()

        p = document.add_paragraph()
        for i, word in enumerate(self.words):
            k = int(255 * (1 - word.conf))
            if i in punctuation:
                if '.' in punctuation[i]:
                    self.words[i + 1].value = self.words[i + 1].value.capitalize()
                run = p.add_run(word.value)
                if k != 0:
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(k, k, k)

                p.add_run(punctuation[i].rstrip('\n') + ' ')

                if '\n' in punctuation[i]:
                    p = document.add_paragraph()
            else:
                run = p.add_run(word.value + ' ')
                if k != 0:
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(k, k, k)

        file_name += '.docx'
        document.save(file_name)

        return file_name
